# SPDX-License-Identifier: AGPL-3.0-only
# Copyright 2026 flxk1
"""Format-aware text extraction for PDF, DOCX, Pages, and other binary docs.

Drop-in replacement for the plain-text-only ``DefaultExtractor``. Reads the
text out of each supported format, then produces the same ``ExtractedFile``
shape — so the downstream classifier and ND router fire on real text
instead of seeing ``(binary file, N bytes, mime=...)`` stubs.

Architecture:

- ``FormatAwareExtractor`` dispatches by file suffix (and falls back on
  mime-type) to a registered text-reader function.
- Each reader returns ``str`` (possibly with newlines / paragraph breaks)
  on success, or raises so the caller can fall back to the metadata stub.
- The dispatch is open — new formats register a new reader function with
  ``register_reader(suffix, reader_fn)``. Don't subclass.

Optional deps (set in ``pyproject.toml`` under ``[project.optional-dependencies]``):

- ``pypdf`` for PDF
- ``python-docx`` for DOCX
- (Pages is unzipped + index.xml parsed with stdlib zipfile + xml.etree)

If an optional dep is missing, that format reader is skipped at runtime
and the extractor falls back to a metadata stub for that file (same
behaviour as today). No hard import-time failure.
"""

from __future__ import annotations

import mimetypes
import time
import zipfile
import xml.etree.ElementTree as ET
from dataclasses import dataclass
from pathlib import Path
from typing import Callable, Optional

from .inbox_watcher import ExtractedFile, _hash_file


# ---------------------------------------------------------------------------
# Per-format text readers
# ---------------------------------------------------------------------------
#
# Each reader: takes a Path, returns a str (possibly with newlines preserving
# paragraph structure), raises on failure. Optional deps are imported lazily
# so the module loads even when pypdf / python-docx aren't installed.


def _read_plain_text(p: Path) -> str:
    """Plain-text formats: .txt, .md, .csv, .json, .xml, .html, …"""
    return p.read_text(encoding="utf-8", errors="replace")


def _read_pdf_pdftotext(p: Path) -> str | None:
    """Layout-aware extraction via poppler's ``pdftotext -layout`` when present. Far better on
    multi-column / justified PDFs (EUR-Lex regulations) where pypdf inserts spurious mid-word
    spaces. Returns None if pdftotext is unavailable or fails → caller falls back to pypdf."""
    import shutil
    import subprocess
    exe = shutil.which("pdftotext")
    if not exe:
        return None
    try:
        out = subprocess.run([exe, "-layout", "-q", str(p), "-"],
                             capture_output=True, timeout=120)
    except (OSError, subprocess.SubprocessError):
        return None
    if out.returncode != 0:
        return None
    text = out.stdout.decode("utf-8", errors="replace")
    return text if text.strip() else None


def _read_pdf(p: Path) -> str:
    """PDF text extraction. Prefers poppler ``pdftotext -layout`` (column-aware); falls back
    to pypdf. Best-effort; skips encrypted PDFs."""
    layout = _read_pdf_pdftotext(p)
    if layout is not None:
        return layout
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise RuntimeError("pypdf not installed (pip install pypdf)") from e

    reader = PdfReader(str(p))
    if reader.is_encrypted:
        # Try empty-password unlock; if that fails, skip.
        try:
            reader.decrypt("")
        except Exception:
            raise RuntimeError("PDF is encrypted; skipping")

    parts: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            txt = page.extract_text() or ""
        except Exception:
            txt = ""
        if txt.strip():
            parts.append(txt)
    # form-feed page breaks — the same contract pdftotext emits, so downstream
    # extraction sees identical text whichever backend produced it
    return "\f".join(parts)


def _read_docx(p: Path) -> str:
    """DOCX text extraction via python-docx. Paragraphs + table cells."""
    try:
        import docx as _docx  # python-docx
    except ImportError as e:
        raise RuntimeError("python-docx not installed (pip install python-docx)") from e

    doc = _docx.Document(str(p))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def _read_legal_xml(p: Path) -> str:
    """EUR-Lex legal XML (Akoma Ntoso / Formex): parse the structure and return
    a block-structured text projection (one provision per block).

    Falls back to raising on non-legal XML so the dispatcher uses the plain-text
    reader for ordinary XML files. A legal act is detected by an
    ``akomaNtoso`` root or an ``ARTICLE`` element anywhere in the tree.
    """
    raw = p.read_bytes()
    head = raw[:4000].lower()
    if b"akomantoso" not in head and b"<article" not in head:
        raise RuntimeError("not a recognised legal XML (no akomaNtoso/ARTICLE)")
    from .xml_legal import parse_legal_xml, document_tree_to_text
    tree = parse_legal_xml(raw)
    if not tree.nodes:
        raise RuntimeError("legal XML parsed but no provisions found")
    return document_tree_to_text(tree)


def _read_pages(p: Path) -> str:
    """Apple Pages (.pages): zip archive containing index.xml.

    Reads index.xml, strips XML, returns the visible text. Best-effort —
    Pages's format is undocumented and varies by version. Skips on failure.
    """
    try:
        with zipfile.ZipFile(str(p)) as zf:
            # Try a few known locations for the document XML.
            candidates = [
                "index.xml",
                "Index.zip",  # newer Pages stores binary; we'd need to handle this differently
                "preview-web.html",
            ]
            for name in candidates:
                if name in zf.namelist():
                    if name.endswith(".xml"):
                        raw = zf.read(name).decode("utf-8", errors="replace")
                        try:
                            root = ET.fromstring(raw)
                            text_parts = [t for t in root.itertext() if t and t.strip()]
                            return " ".join(text_parts)
                        except ET.ParseError:
                            return raw  # best effort: return raw XML
                    if name.endswith(".html"):
                        raw = zf.read(name).decode("utf-8", errors="replace")
                        # Strip HTML tags conservatively.
                        import re as _re
                        return _re.sub(r"<[^>]+>", " ", raw)
            raise RuntimeError(".pages archive has no readable index — likely a binary-format Pages doc")
    except zipfile.BadZipFile as e:
        raise RuntimeError(".pages file is not a valid zip") from e


# ---------------------------------------------------------------------------
# Reader registry
# ---------------------------------------------------------------------------

Reader = Callable[[Path], str]

_DEFAULT_READERS: dict[str, Reader] = {
    # Plain text
    ".txt":   _read_plain_text,
    ".md":    _read_plain_text,
    ".markdown": _read_plain_text,
    ".csv":   _read_plain_text,
    ".tsv":   _read_plain_text,
    ".json":  _read_plain_text,
    # Legal XML (Akoma Ntoso / Formex) parses structure-aware via
    # _read_xml_dispatch, which degrades to plain text for ordinary XML.
    ".xml":   None,   # bound below to _read_xml_dispatch (defined after it)
    ".akn":   None,
    ".html":  _read_plain_text,
    ".htm":   _read_plain_text,
    ".rtf":   _read_plain_text,    # not strictly plain but close enough for now
    ".log":   _read_plain_text,
    # Office formats
    ".pdf":   _read_pdf,
    ".docx":  _read_docx,
    ".pages": _read_pages,
}


def _read_xml_dispatch(p: Path) -> str:
    """`.xml`/`.akn` reader: try structure-aware legal parsing, else plain text.

    A EUR-Lex Akoma Ntoso / Formex act parses into block-structured provisions;
    any other XML degrades to its raw text so nothing is lost.
    """
    try:
        return _read_legal_xml(p)
    except Exception:
        return _read_plain_text(p)


_DEFAULT_READERS[".xml"] = _read_xml_dispatch
_DEFAULT_READERS[".akn"] = _read_xml_dispatch


def register_reader(suffix: str, reader: Reader) -> None:
    """Register or replace the reader for a file suffix (case-insensitive)."""
    _DEFAULT_READERS[suffix.lower()] = reader


# ---------------------------------------------------------------------------
# FormatAwareExtractor
# ---------------------------------------------------------------------------

MAX_TEXT_BYTES = 2 * 1024 * 1024   # 2 MiB cap on extracted text per file


@dataclass
class _ReadResult:
    text: str
    reader_name: str
    truncated: bool
    error: Optional[str] = None


def _extract_text(p: Path) -> _ReadResult:
    """Dispatch by suffix; cap at MAX_TEXT_BYTES; never raise."""
    suffix = p.suffix.lower()
    reader = _DEFAULT_READERS.get(suffix)
    if reader is None:
        return _ReadResult(text="", reader_name="none",
                           truncated=False,
                           error=f"no reader for suffix '{suffix}'")
    try:
        raw = reader(p)
    except Exception as e:
        return _ReadResult(text="", reader_name=reader.__name__,
                           truncated=False,
                           error=f"{type(e).__name__}: {e}")
    if not isinstance(raw, str):
        raw = str(raw)
    encoded = raw.encode("utf-8", errors="replace")
    if len(encoded) > MAX_TEXT_BYTES:
        return _ReadResult(
            text=encoded[:MAX_TEXT_BYTES].decode("utf-8", errors="replace"),
            reader_name=reader.__name__,
            truncated=True,
        )
    return _ReadResult(text=raw, reader_name=reader.__name__, truncated=False)


class FormatAwareExtractor:
    """Extractor that reads text from PDF / DOCX / Pages / plain-text files
    and produces a single pair carrying the extracted text. Downstream
    components (classifier + NDs) can then fire on the actual content.

    Drop-in replacement for ``DefaultExtractor`` in the inbox-watcher path.
    Same idempotency contract (pair_id = file_hash).
    """

    extractor_id = "format-aware-extractor"
    extractor_version = "0.1.0"

    def extract(self, file_path: str, folder_context: str) -> ExtractedFile:
        p = Path(file_path)
        size = p.stat().st_size
        h = _hash_file(p)
        mime, _ = mimetypes.guess_type(str(p))
        mime = mime or "application/octet-stream"

        result = _extract_text(p)
        if result.error and not result.text:
            # No reader, or read failed — fall back to a metadata stub.
            body = f"(binary file, {size} bytes, mime={mime}, extract_error={result.error})"
            body_format = "metadata"
        else:
            body = result.text
            body_format = "prose"

        # Tier D — ingest-time prompt-injection scan (A5 mitigation). Runs
        # before the body can be forwarded into any downstream LLM prompt.
        # Non-blocking: findings are recorded onto the pair (and thus the
        # signed log) so a caller forwarding the body can refuse/quarantine.
        injection_findings: list[dict] = []
        try:
            from workspaces.lock import scan_document as _scan_injection
            injection_findings = [
                {"type": f.type, "severity": f.severity, "detail": f.detail}
                for f in _scan_injection(body)
            ]
        except Exception:
            injection_findings = []

        pair_id = h
        problem_id = f"sha256:problem-{h[7:]}"
        pair: dict = {
            "id": pair_id,
            "problem": {
                "id": problem_id,
                "scope": "inbox",
                "type": "document_ingest",
                "summary": p.name,
                "facets": {
                    "filename": p.name,
                    "size_bytes": size,
                    "mime_type": mime,
                    "suffix": p.suffix.lower(),
                    "reader": result.reader_name,
                    "truncated": result.truncated,
                    "ingested_at": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
                    "injection_findings": injection_findings,
                    "injection_flagged": bool(injection_findings),
                },
                "source_document": str(p.resolve()),
            },
            "solution": {
                "id": pair_id,
                "problem_id": problem_id,
                "body": body,
                "body_format": body_format,
                "authority_tier": 3,
                "confidence": 1.0 if not result.error else 0.5,
                "cited_sources": [str(p.resolve())],
                "extractor_chain": [self.extractor_id, result.reader_name],
                "extractor_version": self.extractor_version,
                "created_at": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
            },
        }
        # content_preview is what the classifier + NDs read. Use the full
        # extracted text (capped at MAX_TEXT_BYTES already) so domain NDs
        # see actual content from multi-page documents, not a 1-KiB preview.
        return ExtractedFile(
            file_path=str(p.resolve()),
            file_size=size,
            file_hash=h,
            mime_type=mime,
            content_preview=body if result.error is None else "",
            pairs=[pair],
        )
